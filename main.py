import fitz
import easyocr
from docx import Document
import re
import torch
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fuzzywuzzy import fuzz

def clear_word_table(word_path):
    doc = Document(word_path)
    table = doc.tables[0]
    for row in table.rows[2:]:  
        table._element.remove(row._element)  # Удаляем строку
    doc.save(word_path)

def process_pdf(pdf_path, keywords, word_path, threshold, languages, text_q, count):
    count = 1

    reader = easyocr.Reader(['en', languages], gpu=True)

    # Поиск ключевых слов и даты
    found_keywords = []  # Изменение: используем список для хранения ключевых слов
    found_date = None
    found_outgoing_num = None
    found_outgoing_num2 = None
    found_outgoing_num_cut = None
    found_outgoing_num2_cut = None
    total_text = ""
    start_page = 1

    with fitz.open(pdf_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf.load_page(page_num)
            text = page.get_text()

            print(f"Обрабатывается страница {page_num + 1}...")

            if text_q == 'y':
                print("Распознанный текст на странице с ручного ввода: \n")
                print(text, '\n')

            if text:
                total_text += text

                # Поиск ключевых слов
                for keyword in keywords:
                    if keyword in found_keywords:
                        continue
                    key_words = keyword.split()
                    found_count = sum(word in total_text for word in key_words)
                    recognition_percentage = (found_count / len(key_words)) * 100
                    if recognition_percentage >= threshold and all(word in total_text for word in key_words):
                        print(f"Ключевое слово '{keyword}' добавлено с процентом распознавания {recognition_percentage}%")
                        found_keywords.append(keyword)
                    else:
                        print(f"Ключевое слово '{keyword}' не добавлено с процентом распознавания {recognition_percentage}%")

                # Поиск исходящего номера, если он еще не был найден
                if not found_outgoing_num:
                    outgoing_num = find_first_matching_number(text)
                    if outgoing_num:
                        print(f"Найден исходящий номер {outgoing_num}")
                        found_outgoing_num = outgoing_num
                        if len(found_outgoing_num) >= 10:
                            found_outgoing_num = outgoing_num[:9]
                            found_outgoing_num_cut = outgoing_num[9:]
                            print(f"Исходящий номер разделен на две части: {found_outgoing_num}, {found_outgoing_num_cut}")

                if not found_outgoing_num2:
                    outgoing_num2 = find_first_matching_number2(text)
                    if outgoing_num2:
                        print(f"Найден входящий номер {outgoing_num2}")
                        found_outgoing_num2 = outgoing_num2
                        if len(found_outgoing_num2) >= 10:
                            found_outgoing_num2 = outgoing_num2[:9]
                            found_outgoing_num2_cut = outgoing_num2[9:]
                            print(f"Исходящий номер разделен на две части: {found_outgoing_num2}, {found_outgoing_num2_cut}")

            # Проверка изображений на странице
            images = page.get_images(full=True)
            if images:
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    result = reader.readtext(image_bytes)
                    for detection in result:
                        img_text = detection[1]
                        total_text += " " + img_text

                for keyword in keywords:
                    if keyword in found_keywords:
                        continue
                    key_words = keyword.split()
                    found_count = sum(word in total_text for word in key_words)
                    recognition_percentage = (found_count / len(key_words)) * 100
                    if recognition_percentage >= threshold and all(word in total_text for word in key_words):
                        print(f"Ключевое слово '{keyword}' добавлено с процентом распознавания {recognition_percentage}%")
                        found_keywords.append(keyword)
                    else:
                        print(f"Ключевое слово '{keyword}' не добавлено с процентом распознавания {recognition_percentage}%")

                # Поиск исходящего номера, если он еще не был найден
                if not found_outgoing_num:
                    outgoing_num = find_first_matching_number(text)
                    if outgoing_num:
                        print(f"Найден исходящий номер {outgoing_num}")
                        found_outgoing_num = outgoing_num
                        if len(found_outgoing_num) >= 10:
                            found_outgoing_num = outgoing_num[:9]
                            found_outgoing_num_cut = outgoing_num[9:]
                            print(f"Исходящий номер разделен на две части: {found_outgoing_num}, {found_outgoing_num_cut}")

                if not found_outgoing_num2:
                    outgoing_num2 = find_first_matching_number2(text)
                    if outgoing_num2:
                        print(f"Найден входящий номер {outgoing_num2}")
                        found_outgoing_num2 = outgoing_num2
                        if len(found_outgoing_num2) >= 10:
                            found_outgoing_num2 = outgoing_num2[:9]
                            found_outgoing_num2_cut = outgoing_num2[9:]
                            print(f"Исходящий номер разделен на две части: {found_outgoing_num2}, {found_outgoing_num2_cut}")

            if text_q == 'y':
                print("Распознанный текст на странице с изображения:\n")
                print(total_text, '\n')

            if text_q == 'y':
                # Вывод корректно опознанных слов из ключа
                for found_keyword in found_keywords:
                    key_words_list = found_keyword.split()
                    correct_words = [word for word in key_words_list if word in total_text]
                    if correct_words:
                        print(f"Корректно опознанные слова из ключа '{found_keyword}': {', '.join(correct_words)}")

            if len(pdf) == 1:
                print("Документ содержит только одну страницу. Завершение обработки.\n")
                end_page = 1
                update_word_table(word_path, keywords, found_keywords, found_date, start_page, end_page,
                                   found_outgoing_num, found_outgoing_num2, count, found_outgoing_num2_cut,
                                   found_outgoing_num_cut)
                found_keywords = []
                found_date = None
                found_outgoing_num = None
                found_outgoing_num2 = None
                found_outgoing_num_cut = None
                found_outgoing_num2_cut = None
            else:
                if "End" in text:
                    print(f"Найдена пометка 'End' на странице {page_num + 1}. Завершение документа.")
                    end_page = page_num + 1
                    update_word_table(word_path, keywords, found_keywords, found_date, start_page, end_page,
                                       found_outgoing_num, found_outgoing_num2, count, found_outgoing_num2_cut,
                                       found_outgoing_num_cut)

                    count += 1
                    found_keywords = []
                    found_date = None
                    found_outgoing_num = None
                    found_outgoing_num2 = None
                    found_outgoing_num_cut = None
                    found_outgoing_num2_cut = None
                    total_text = ""

                    start_page = page_num + 2

    return found_keywords, found_date





def update_word_table(word_path, keywords, found_keywords, found_date, start_page, end_page, found_outgoing_num, found_outgoing_num2, count, found_outgoing_num2_cut, found_outgoing_num_cut):
    doc = Document(word_path)
    table = doc.tables[0]
    outgoing_index = None  # Добавляем индекс для столбца с исходящим номером
    null_row = None
    new_row1 = None
    emty_row2 = None
    emty_row = None

    font_name = 'Times New Roman'
    font_size = Pt(12)
    # Находим индекс столбца "Наименование документа"
    for cell in table.rows[0].cells:
        if cell.text.strip() == "Наименование документа":
            column_index = cell._element.getparent().index(cell._element)
            break
    for cell in table.rows[0].cells:
        if cell.text.strip() == "Номера листов":
            list_index = cell._element.getparent().index(cell._element)
            break   
    for cell in table.rows[1].cells:
        if cell.text.strip() == "исходящий":
            outgoing_index = cell._element.getparent().index(cell._element) - 1
            break 
    for cell in table.rows[0].cells:
        if cell.text.strip() == "№ з/п":
            num_index = cell._element.getparent().index(cell._element) - 2
            break 
    for cell in table.rows[1].cells:
        if cell.text.strip() == "входящий":
            ingoing_index = cell._element.getparent().index(cell._element) - 1
            break 

    # Добавляем новую строку в таблицу
    new_row_index = len(table.rows)
    new_row = table.add_row()
    added_keywords = []
    # Если найдены ключевые слова, добавляем их и дату в таблицу
    cell = table.cell(new_row_index, column_index)
    if found_keywords:
        found_keyword = found_keywords[0]  # Берем только первое найденное ключевое слово
        key_description = keywords.get(found_keyword)

        if key_description is None:
            print(f"Описание для ключа '{found_keyword}' не найдено.")
            return
        
        key_text = key_description['description']
        key_text2 = key_description['description2'] # Получаем description2, если он есть, или пустую строку
        key_text3 = key_description['description3'] # Получаем description3, если он есть, или пустую строку
        key_text4 = key_description['description4'] # Получаем description4, если он есть, или пустую строку
        key_text5 = key_description['description5'] # Получаем description5, если он есть, или пустую строку


    
        if key_text2 != "":
            if found_date:
                key_text2 += f", от {found_date}"
            is_two_str = True
            new_row1 = table.add_row()
            column_cell = new_row1.cells[column_index]  # Получаем ячейку в нужном столбце
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем выравнивание по левому краю
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Используем column_cell вместо column_index
            column_cell.text = key_text2
            for paragraph in column_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size

        if key_text3 != "":
            if found_date:
                key_text3 += f", от {found_date}"
            is_two_str = True
            new_row = table.add_row()
            column_cell = new_row.cells[column_index]  # Получаем ячейку в нужном столбце
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем выравнивание по левому краю
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Используем column_cell вместо column_index
            column_cell.text = key_text3
            for paragraph in column_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size
        if key_text4 != "":
            if found_date:
                key_text4 += f", от {found_date}"
            new_row = table.add_row()
            column_cell = new_row.cells[column_index]  # Получаем ячейку в нужном столбце
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем выравнивание по левому краю
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Используем column_cell вместо column_index
            column_cell.text = key_text4
            for paragraph in column_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size
        
        if key_text5 != "":
            if found_date:
                key_text5 += f", от {found_date}"
            new_row = table.add_row()
            column_cell = new_row.cells[column_index]  # Получаем ячейку в нужном столбце
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Устанавливаем выравнивание по левому краю
            column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Используем column_cell вместо column_index
            column_cell.text = key_text5
            for paragraph in column_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size


        else:
            if found_date:
                key_text += f", от {found_date}"
        cell_paragraphs = cell.paragraphs
        if not cell_paragraphs:  # Если в ячейке нет абзацев, создаем новый
            new_paragraph = cell.add_paragraph()
        else:
            new_paragraph = cell_paragraphs[-1]  # Или берем последний абзац, если он уже существует
        # Добавляем текст с форматированием
        new_run = new_paragraph.add_run(key_text)
        new_run.font.name = font_name
        new_run.font.size = font_size


    else:
        null_row = table.add_row()
        column_cell = null_row.cells[column_index]  # Получаем ячейку в нужном столбце
        # Используем column_cell вместо column_index
        null_row.text = ""  # Обновляем текст ячейки с key_text
        # Устанавливаем выравнивание по центру
        column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Устанавливаем выравнивание по левому краю
        column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Добавляем диапазон страниц в столбец "Номера листов"
    if start_page == end_page:
        pages_range = f"{start_page}"" "
    else:
        pages_range = f"{start_page}-{end_page}"

    list_cell = table.cell(new_row_index, list_index)
    list_cell.text = pages_range
    list_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
    # Применяем параметры шрифта к тексту ячейки
    for paragraph in list_cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
    
    # if is_two_str == False:
        # Добавляем номер заказа в соответствующую ячейку
    if outgoing_index is not None:  
        outgoing_cell = table.cell(new_row_index, outgoing_index)
        if found_outgoing_num:
            outgoing_cell.text = found_outgoing_num
            if found_outgoing_num_cut:
                outgoing_cell.text = found_outgoing_num + '-'
                if null_row:
                    column_cell = null_row.cells[1]
                    column_cell.text = found_outgoing_num_cut               
                if new_row1:
                    column_cell = new_row1.cells[1]
                    column_cell.text = found_outgoing_num_cut
                if new_row1 == None and null_row == None:
                    emty_row = table.add_row()
                    column_cell = emty_row.cells[1]
                    column_cell.text = found_outgoing_num_cut
                column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for paragraph in column_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size


        # Устанавливаем выравнивание по центру
        outgoing_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Устанавливаем выравнивание по левому краю
        outgoing_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Применяем параметры шрифта к тексту ячейки
        for paragraph in outgoing_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = font_size
    else:
        print("Столбец 'исходящие' не найден в таблице.")

    if ingoing_index is not None:
        ingoing_index = table.cell(new_row_index, ingoing_index)
        if found_outgoing_num2:
            ingoing_index.text = found_outgoing_num2
            if found_outgoing_num2_cut:
                ingoing_index.text = found_outgoing_num2 + '-'
                if null_row:
                    column_cell = null_row.cells[2]
                    column_cell.text = found_outgoing_num2_cut
                if new_row1:
                    column_cell = new_row1.cells[2]
                    column_cell.text = found_outgoing_num2_cut
                if emty_row:
                    column_cell = emty_row.cells[2]
                    column_cell.text = found_outgoing_num2_cut
                if new_row1 == None and null_row == None and emty_row == None:
                    emty_row2 = table.add_row()
                    column_cell = emty_row2.cells[2]
                    column_cell.text = found_outgoing_num2_cut
                column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                column_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for paragraph in column_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size


        ingoing_index.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        ingoing_index.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for paragraph in ingoing_index.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = font_size
    else:
        print("Столбец 'входящие' не найден в таблице.")


    list_num = table.cell(new_row_index, num_index + 1)
    list_num.text = str(count),'.'
    list_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
    # Применяем параметры шрифта к тексту ячейки
    for paragraph in list_num.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size

    doc.save(word_path)



def find_first_matching_number(text):
    pattern = r'ИСХ№(\d{1,5}(?:\d*[-/]\d*)*дск)'  # Используем группировку для всего номера, кроме "ИСХ№"
    match = re.search(pattern, text)
    if match:
        return match.group(1)  # Возвращаем содержимое первой группы (найденный номер)
    else:
        return None
    

def find_first_matching_number2(text):
    pattern2 = r'ВХ№(\d{1,5}(?:\d*[-/]\d*)*дск)'  # Включаем номер и "дск" в группы, чтобы их можно было извлечь
    match2 = re.search(pattern2, text)
    if match2:
        return match2.group(1)  # Возвращаем значение номера и "дск"
    else:           
        return None


# def find_dates(text):
#     # Шаблон для поиска даты в формате DD.MM.YYYY
#     date_pattern = r'\b\d{2}[,.]?\d{2}[,.]?\d{4}\b'
#     # Находим первое совпадение с шаблоном
#     match = re.search(date_pattern, text)
#     if match:
#         return match.group()
#     else:
#         return None

def read_keys(keys_path):
    keys = {}
    doc = Document(keys_path)
    key = ''  # Переменная для хранения текущего ключа
    description = ''  # Переменная для хранения описания текущего ключа
    description2 = ''  # Переменная для хранения второго описания текущего ключа
    description3 = ''  # Переменная для хранения третьего описания текущего ключа
    description4 = ''  # Переменная для хранения четвертого описания текущего ключа
    description5 = ''  # Переменная для хранения пятого описания текущего ключа
    description6 = ''  # Переменная для хранения шестого описания текущего ключа
    cell_format = None  # Переменная для хранения форматирования текущего ключа
    
    for row in doc.tables[0].rows[1:]:  # Пропускаем первую строку, так как это заголовок
        cell_0_text = row.cells[1].text.strip()
        cell_0_number = row.cells[0].text.strip()

        if cell_0_number:  # Если ячейка не пустая, это начало нового ключа
            # Если есть предыдущий ключ, сохраняем его в словарь
            if key:
                keys[key] = {'description': description, 'description2': description2, 
                             'description3': description3, 'description4': description4,
                             'description5': description5, 'format': cell_format}  # Сохраняем информацию о форматировании в keys
                print(f"Добавлен Ключ: {key}. с описанием: {description} {description2} {description3} {description4} {description5}")
            key = cell_0_text  # Обновляем текущий ключ
            description = row.cells[2].text.strip()  # Берем текст из второй ячейки в строке (столбец "Описание ключа")
            description2 = ''  # Обнуляем description2 для нового ключа
            description3 = ''  # Обнуляем description3 для нового ключа
            description4 = ''  # Обнуляем description4 для нового ключа
            description5 = ''  # Обнуляем description5 для нового ключа
            description6 = ''
            cell_format = {}  # Сбрасываем форматирование для нового ключа
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    cell_format['bold'] = run.bold
                    cell_format['italic'] = run.italic
                    cell_format['underline'] = run.underline
                    cell_format['font_color'] = run.font.color.rgb
                    cell_format['font_size'] = run.font.size
                    cell_format['font_name'] = run.font.name
                    cell_format['highlight_color'] = run.font.highlight_color
                    cell_format['superscript'] = run.font.superscript
                    cell_format['subscript'] = run.font.subscript
                    cell_format['strike'] = run.font.strike
                    cell_format['double_strike'] = run.font.double_strike
                    cell_format['all_caps'] = run.font.all_caps
                    cell_format['small_caps'] = run.font.small_caps
                    cell_format['shadow'] = run.font.shadow
                    cell_format['outline'] = run.font.outline
                    cell_format['emboss'] = run.font.emboss
                    cell_format['imprint'] = run.font.imprint
        else:
            # Если ячейка пустая, это продолжение описания или ключа
            if not description2:
                description2 = row.cells[2].text.strip()
            elif not description3:
                description3 = row.cells[2].text.strip()
            elif not description4:
                description4 = row.cells[2].text.strip()
            elif not description5:
                description5 = row.cells[2].text.strip()
            elif not description6:
                description6 = row.cells[2].text.strip()
            else:
                key += " " + row.cells[1].text.strip()  # Добавляем новую строку к текущему ключу

    # Сохраняем информацию о последнем ключе
    if key:
        keys[key] = {'description': description, 'description2': description2, 
                     'description3': description3, 'description4': description4,
                     'description5': description5, 'description6': description6,'format': cell_format}  # Сохраняем информацию о форматировании в keys
        print(f"Добавлен Ключ: {key}. с описанием: {description} {description2} {description3} {description4} {description5} {description6}")

    return keys

if __name__ == "__main__":
    file_path = input("Введите путь к файлу (PDF): ")
    threshold = int(input("Введите минимальное пороговое значение для распознавания текста в %: "))
    languages = input("Введите язык для использования (ru или uk или be) ")
    text_q = input("Выводить распознаный текст на страницах и слова из каждого ключа? (y = да, n = нет) ")

    word_path = "result.docx"
    keys_path = "keys.docx"
    keywords = read_keys(keys_path)
    clear_word_table(word_path)  # Очищаем таблицу перед обработкой нового файла PDF
    print("Таблица 'Result.docx' очищена \n")
    count = 1
    found_keywords, found_date = process_pdf(file_path, keywords, word_path, threshold, languages, text_q, count)
    try:
        update_word_table(word_path, keywords, found_keywords, found_date, count)  # Передаем словарь с описаниями ключей в функцию
    except Exception as e:
        print("Конец")

    input("Press Enter to exit...")
